"""Convert DOCX to markdown text. Images are replaced with a placeholder."""
import pathlib
from docx import Document
from docx.oxml.ns import qn


def _has_image(paragraph) -> bool:
    """Return True if paragraph contains an inline image."""
    for run in paragraph.runs:
        if run._element.findall('.//' + qn('a:blip')):
            return True
    return False


def _extract_tekstvakken(doc) -> list[str]:
    """Verzamel tekst uit text frames / tekstvakken (w:txbxContent).

    python-docx doorloopt standaard alleen hoofd-paragrafen en tabellen —
    tekst in text frames (drawings, shapes, callouts) wordt gemist. Voor
    gemeentelijke documenten bevatten tekstvakken vaak colofons met
    contactgegevens (adres, telefoon, e-mail).
    """
    teksten: list[str] = []
    # Zoek in document body, headers en footers
    bronnen = [doc.element.body]
    for section in doc.sections:
        if section.header is not None:
            bronnen.append(section.header._element)
        if section.footer is not None:
            bronnen.append(section.footer._element)

    for bron in bronnen:
        for txbx in bron.iter(qn("w:txbxContent")):
            # Verzamel alle w:t elementen binnen dit tekstvak
            onderdelen = [t.text or "" for t in txbx.iter(qn("w:t"))]
            tekst = "".join(onderdelen).strip()
            if tekst:
                teksten.append(tekst)
    return teksten


def _extract_header_footer_tekst(doc) -> list[str]:
    """Verzamel hoofdtekst (buiten tekstvakken) uit headers en footers."""
    teksten: list[str] = []
    for section in doc.sections:
        for container in (section.header, section.footer):
            if container is None:
                continue
            for para in container.paragraphs:
                t = para.text.strip()
                if t:
                    teksten.append(t)
    return teksten


def _heading_level(paragraph) -> int:
    """Return heading level 1-6, or 0 if not a heading."""
    name = paragraph.style.name or ""
    if name.startswith("Heading "):
        try:
            return int(name.split(" ")[1])
        except (IndexError, ValueError):
            pass
    return 0


def _table_to_markdown(table) -> str:
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)


def docx_to_markdown(pad: pathlib.Path) -> str:
    doc = Document(pad)
    parts = []

    # Track which paragraphs are inside tables to avoid double-processing
    table_paragraphs = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    table_paragraphs.add(id(para))

    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

        if tag == "tbl":
            # Find matching table object
            for table in doc.tables:
                if table._tbl is block:
                    parts.append(_table_to_markdown(table))
                    break

        elif tag == "p":
            # Find matching paragraph object
            from docx.text.paragraph import Paragraph
            para = Paragraph(block, doc)

            if id(para) in table_paragraphs:
                continue

            if _has_image(para):
                parts.append("[AFBEELDING VERWIJDERD]")
                continue

            tekst = para.text.strip()
            if not tekst:
                continue

            level = _heading_level(para)
            if level:
                parts.append("#" * level + " " + tekst)
            else:
                parts.append(tekst)

    # Voeg tekst uit headers/footers en tekstvakken toe.
    # Deze worden door python-docx niet meegenomen in doc.element.body-iteratie,
    # maar bevatten in gemeentelijke documenten vaak contactgegevens.
    extra_blokken = _extract_header_footer_tekst(doc) + _extract_tekstvakken(doc)
    if extra_blokken:
        parts.append("")  # scheiding
        parts.extend(extra_blokken)

    return "\n\n".join(parts)
